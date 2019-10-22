import xlsxwriter as xls
#workbook=xls.Workbook('Template.xlsx')
#worksheet=workbook.add_worksheet('FMEA')

expenses =(
    ['rent',100],
    ['gas',100],
)

row=0
col=0
for item, cost in expenses:
    #worksheet.write(row,col, item)
    #worksheet.write(row,col+1, cost)
    row+=1


#workbook.close()

import pyexcel as pe
book=pe.get_book(file_name="Template.xlsx")
print(book.sheet_names())
sheet=book.sheet_by_name(book.sheet_names()[1])
print(sheet)

#print(sheet[1,1])



fp = open('styles.txt')
styles=fp.readlines()
styles=[x.strip() for x in styles]
#styles=[x.replace(' ','') for x in styles]

from docx import Document

document = Document()
table = document.add_table(rows=3,cols=3)
table.style = 'Colorful Grid Accent 1'
cell=table.cell(0,1)
cell.text=sheet[1,1]
print(table.style)

for style in styles:
    table = document.add_table(rows=3, cols=3)
    table.style = style
    cell = table.cell(0, 1)
    cell.text = style
    print(table.style)


document.save('demo.docx')

